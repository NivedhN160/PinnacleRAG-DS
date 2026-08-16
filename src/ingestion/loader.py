"""
Multi-format document loader for PinnacleRAG-DS.

Supports PDF, TXT, MD, and DOCX files with rich metadata extraction.
"""

import os
from pathlib import Path
from typing import Optional

from langchain_core.documents import Document

from config.settings import Settings
from src.utils.logging import get_logger

logger = get_logger(__name__)


class DocumentLoader:
    """Load documents from various file formats with metadata."""

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}

    def __init__(self, settings: Settings) -> None:
        self.settings = settings

    def load_directory(self, path: str) -> list[Document]:
        """
        Load all supported documents from a directory recursively.
        Detects domain from subfolder structure: data/raw/{domain}/...
        """
        dir_path = Path(path)
        if not dir_path.exists():
            logger.warning(f"Directory not found: {path}")
            return []

        documents: list[Document] = []
        files_found = 0

        for file_path in sorted(dir_path.rglob("*")):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                files_found += 1
                try:
                    docs = self.load_file(str(file_path))
                    documents.extend(docs)
                except Exception as e:
                    logger.error(f"Failed to load {file_path}: {e}")

        logger.info(f"Loaded {len(documents)} document(s) from {files_found} file(s) in {path}")
        return documents

    def _detect_domain(self, file_path: Path) -> str:
        """Detect domain from file path based on data/raw/{domain}/ structure."""
        raw_path = Path(self.settings.raw_data_path)
        try:
            relative = file_path.relative_to(raw_path)
            parts = relative.parts
            if len(parts) > 1:
                domain = parts[0].lower()
                if domain in ("trading", "security", "seo", "general"):
                    return domain
        except (ValueError, IndexError):
            pass
        return "general"

    def load_file(self, file_path: str) -> list[Document]:
        """Load a single file with metadata including domain."""
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {suffix}. Supported: {self.SUPPORTED_EXTENSIONS}")

        domain = self._detect_domain(path)

        base_metadata = {
            "source": str(path),
            "filename": path.name,
            "file_type": suffix,
            "domain": domain,
        }

        if suffix == ".pdf":
            return self._load_pdf(path, base_metadata)
        elif suffix == ".docx":
            return self._load_docx(path, base_metadata)
        elif suffix in (".txt", ".md"):
            return self._load_text(path, base_metadata)
        else:
            return []

    def supported_extensions(self) -> list[str]:
        """Return list of supported file extensions."""
        return sorted(self.SUPPORTED_EXTENSIONS)

    def _load_pdf(self, path: Path, base_metadata: dict) -> list[Document]:
        """Load a PDF file, one Document per page."""
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        documents = []

        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                metadata = {**base_metadata, "page": page_num, "total_pages": len(reader.pages)}
                documents.append(Document(page_content=text, metadata=metadata))

        logger.debug(f"PDF loaded: {path.name} ({len(documents)} pages with text)")
        return documents

    def _load_docx(self, path: Path, base_metadata: dict) -> list[Document]:
        """Load a DOCX file as a single Document."""
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)

        if not text.strip():
            return []

        metadata = {**base_metadata, "paragraph_count": len(paragraphs)}
        return [Document(page_content=text, metadata=metadata)]

    def _load_text(self, path: Path, base_metadata: dict) -> list[Document]:
        """Load a text or markdown file as a single Document."""
        text = path.read_text(encoding="utf-8", errors="replace")

        if not text.strip():
            return []

        line_count = len(text.splitlines())
        metadata = {**base_metadata, "line_count": line_count}
        return [Document(page_content=text, metadata=metadata)]
